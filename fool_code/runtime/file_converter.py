"""Pluggable file-to-markdown converters.

Architecture:
  - Each converter implements ``FileConverter`` and is registered in ``CONVERTER_REGISTRY``.
  - To add a new format, create a class and register its extensions.
  - ``process_file()`` is the high-level entry point used by the API layer.
"""

from __future__ import annotations

import logging
import shutil
import uuid as _uuid
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import ClassVar

from fool_code.runtime.config import app_data_root

logger = logging.getLogger(__name__)


@dataclass
class ConversionResult:
    """Result of converting a file to markdown."""

    file_id: str
    original_name: str
    cached_path: str
    markdown_path: str
    markdown_text: str
    category: str
    size: int
    preview: str
    meta: dict = field(default_factory=dict)


class FileConverter(ABC):
    """Base class — one per file format family."""

    extensions: ClassVar[set[str]] = set()

    @abstractmethod
    def convert(self, source: Path) -> tuple[str, dict]:
        """Return (markdown_text, meta_dict).

        ``meta_dict`` can carry format-specific info (pages, sheets, word_count …).
        """
        ...


class DocxConverter(FileConverter):
    extensions = {".docx", ".doc"}

    def convert(self, source: Path) -> tuple[str, dict]:
        from docx import Document

        doc = Document(str(source))
        parts: list[str] = []
        word_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append("")
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            elif "heading" in style_name:
                parts.append(f"#### {text}")
            else:
                parts.append(text)

            word_count += len(text.split())

        for table in doc.tables:
            parts.append("")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
                if ri == 0:
                    parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
                word_count += sum(len(c.split()) for c in cells)

        md = "\n".join(parts).strip()
        meta = {
            "word_count": word_count,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        return md, meta


class XlsxConverter(FileConverter):
    extensions = {".xlsx", ".xls"}

    def convert(self, source: Path) -> tuple[str, dict]:
        from openpyxl import load_workbook

        wb = load_workbook(str(source), read_only=True, data_only=True)
        parts: list[str] = []
        total_rows = 0
        sheet_names: list[str] = []

        for ws in wb.worksheets:
            sheet_names.append(ws.title)
            parts.append(f"## Sheet: {ws.title}")
            parts.append("")

            rows_data: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows_data.append(cells)

            if not rows_data:
                parts.append("*(empty sheet)*")
                continue

            max_cols = max(len(r) for r in rows_data)
            for r in rows_data:
                r.extend([""] * (max_cols - len(r)))

            for ri, cells in enumerate(rows_data):
                safe = [c.replace("|", "\\|").replace("\n", " ") for c in cells]
                parts.append("| " + " | ".join(safe) + " |")
                if ri == 0:
                    parts.append("| " + " | ".join(["---"] * max_cols) + " |")

            total_rows += len(rows_data)
            parts.append("")

        wb.close()
        md = "\n".join(parts).strip()
        meta = {
            "sheets": sheet_names,
            "sheet_count": len(sheet_names),
            "total_rows": total_rows,
        }
        return md, meta


class CsvConverter(FileConverter):
    extensions = {".csv", ".tsv"}

    def convert(self, source: Path) -> tuple[str, dict]:
        import csv

        delimiter = "\t" if source.suffix.lower() == ".tsv" else ","
        text = source.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(text.splitlines(), delimiter=delimiter)
        rows = list(reader)

        if not rows:
            return "*(empty file)*", {"total_rows": 0}

        parts: list[str] = []
        max_cols = max(len(r) for r in rows) if rows else 0
        for ri, cells in enumerate(rows):
            cells.extend([""] * (max_cols - len(cells)))
            safe = [c.replace("|", "\\|").replace("\n", " ") for c in cells]
            parts.append("| " + " | ".join(safe) + " |")
            if ri == 0:
                parts.append("| " + " | ".join(["---"] * max_cols) + " |")

        md = "\n".join(parts)
        return md, {"total_rows": len(rows)}


class PlainTextConverter(FileConverter):
    extensions = {".txt", ".md", ".log", ".json", ".xml", ".yaml", ".yml", ".toml"}

    def convert(self, source: Path) -> tuple[str, dict]:
        text = source.read_text(encoding="utf-8", errors="replace")
        lines = text.splitlines()
        return text, {"lines": len(lines), "word_count": len(text.split())}


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

CONVERTER_REGISTRY: dict[str, FileConverter] = {}


def _register(converter: FileConverter) -> None:
    for ext in converter.extensions:
        CONVERTER_REGISTRY[ext.lower()] = converter


_register(DocxConverter())
_register(XlsxConverter())
_register(CsvConverter())
_register(PlainTextConverter())


def get_converter(path: Path) -> FileConverter | None:
    ext = path.suffix.lower()
    return CONVERTER_REGISTRY.get(ext)


# ---------------------------------------------------------------------------
# High-level entry point
# ---------------------------------------------------------------------------

_CATEGORY_MAP: dict[str, str] = {}
for _conv in (DocxConverter, XlsxConverter, CsvConverter, PlainTextConverter):
    _cat = {
        DocxConverter: "document",
        XlsxConverter: "spreadsheet",
        CsvConverter: "spreadsheet",
        PlainTextConverter: "text",
    }[_conv]
    for _ext in _conv.extensions:
        _CATEGORY_MAP[_ext.lower()] = _cat


def file_cache_dir(session_id: str) -> Path:
    d = app_data_root() / "file-cache" / session_id
    d.mkdir(parents=True, exist_ok=True)
    return d


import re as _re

_MD_HEADING = _re.compile(r"^#{1,6}\s+", _re.MULTILINE)
_MD_TABLE_SEP = _re.compile(r"^\|[\s\-:|]+\|$", _re.MULTILINE)
_MD_PIPE = _re.compile(r"\s*\|\s*")
_MD_EMPHASIS = _re.compile(r"\*{1,2}(.+?)\*{1,2}")


def _clean_preview(md_text: str, max_len: int = 200) -> str:
    """Strip markdown syntax to produce a clean plain-text preview."""
    text = _MD_TABLE_SEP.sub("", md_text)
    text = _MD_HEADING.sub("", text)
    text = _MD_EMPHASIS.sub(r"\1", text)
    text = text.replace("\\|", "|")
    lines = [ln.strip().strip("|").strip() for ln in text.splitlines() if ln.strip()]
    joined = "  ".join(lines)
    if len(joined) > max_len:
        return joined[:max_len] + "…"
    return joined


def process_file(source_path: str, session_id: str) -> ConversionResult | None:
    """Copy file to cache, convert to markdown, return result.

    Returns None if the file type is not supported or conversion fails.
    """
    src = Path(source_path).resolve()
    if not src.is_file():
        logger.warning("File not found: %s", src)
        return None

    converter = get_converter(src)
    if converter is None:
        logger.info("No converter for %s", src.suffix)
        return None

    file_id = f"file-{_uuid.uuid4().hex[:12]}"
    cache_dir = file_cache_dir(session_id)

    safe_name = src.name.replace(" ", "_")
    cached = cache_dir / f"{file_id}_{safe_name}"
    shutil.copy2(str(src), str(cached))

    try:
        md_text, meta = converter.convert(cached)
    except Exception as exc:
        logger.error("Conversion failed for %s: %s", src.name, exc)
        return None

    md_path = cached.with_suffix(cached.suffix + ".md")
    md_path.write_text(md_text, encoding="utf-8")

    preview = _clean_preview(md_text, max_len=200)

    category = _CATEGORY_MAP.get(src.suffix.lower(), "other")

    return ConversionResult(
        file_id=file_id,
        original_name=src.name,
        cached_path=str(cached),
        markdown_path=str(md_path),
        markdown_text=md_text,
        category=category,
        size=src.stat().st_size,
        preview=preview,
        meta=meta,
    )
